from docx import Document


def clean_text(text):
    cleaned_text = ' '.join(text.split())
    return cleaned_text


def lugat1():
    # conn = sqlite3.connect('unicon.db')
    # cursor = conn.cursor()
    file = "L 1 Краткий русско-узбекский словарь по телекоммуникациям.docx"
    doc = Document(docx=file)
    r = 0
    frond_word = ''
    for para in doc.paragraphs:
        if " — " in para.text:
            r += 1
            words = para.text.split(" — ")
            if " " in words[0][0]:
                word_ru = (frond_word + ' ' + words[0].strip()).replace('  ', ' ')
                word_uz_kr = words[1].strip().replace('  ', ' ')
            else:
                word_ru = words[0].strip().replace('  ', ' ')
                word_uz_kr = words[1].strip().replace('  ', ' ')
                frond_word = word_ru.replace('  ', ' ').replace(' (', '(')
            print(word_ru)
            # cursor.execute(f'''
            #             INSERT INTO content_analyze (word_uz_kr, word_ru, soha_id) VALUES
            #             ('{word_uz_kr}', '{word_ru}', 1)''')
    # conn.commit()
    print(r)


def lugat2():
    file = "L 2 Словарь терминов по метрологии в сфере связи и информатизации.docx"
    doc = Document(docx=file)
    # conn = sqlite3.connect('unicon.db')
    # cursor = conn.cursor()

    tables = doc.tables
    q = 0
    for table in tables:
        for row in table.rows:
            if "Атама" in row.cells[1].text:
                print("galdi getti")
                continue
            try:
                # tarifini olish
                words = row.cells[1].text.strip()
                words_ru_index = words.find("ru - ")
                words_en_index = words.find("en - ")

                word_uz_kr = words[:words_ru_index].strip().replace('\n', ' ').replace('  ', ' ')
                word_ru = words[words_ru_index:words_en_index].strip().replace('\n', ' ').replace('  ', ' ')[5:]
                word_en = words[words_en_index:].strip().replace('\n', ' ').replace('  ', ' ')[5:]
                desc = row.cells[2].text.strip()

                # cursor.execute(f'''
                # INSERT INTO content_analyze (word_uz_kr, word_en, word_ru, desc_uz_kr, soha_id)
                # VALUES ('{word_uz_kr}', '{word_en}', '{word_ru}', '{desc}', 2)''')
                q += 1
            except:
                pass
    # conn.commit()
    print(q)
lugat2()

def lugat3():
    file = "L 3 Русско-узбекский толковый словарь.docx"
    doc = Document(docx=file)

    dont = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ё', 'Ж', 'З', 'И', 'Й', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У',
            'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ъ', 'Ы', 'Ь', 'Э', 'Ю', 'Я']

    tables = doc.tables
    # first_tables = tables[0]
    # previous_row = ''
    # previous_text = ''
    # krill_bool = False
    # for row in first_tables.rows:
    #     if row.cells[0]:
    #         row_text = row.cells[0].text.strip()
    #         desc_text = row.cells[1].text.strip()
    #     else:
    #         row_text = previous_row
    #         desc_text = previous_text+' '+row.cells[1].text.strip()
    #     if desc_text:
    #         for krill in dont:
    #             if krill in row_text:
    #                 krill_bool = True
    #         desk_lt = ''
    #         desk_kr = ''
    #         desk_en = ''
    #         desk_ru = ''
    #         if krill_bool:
    #             if len(desc_text.split(' / ')) == 3:
    #                 desk_kr = desc_text.split(' / ')[1]
    #                 desk_en = desc_text.split(' / ')[2]
    #                 desk_ru = desc_text.split(' / ')[0]
    #             else:
    #                 desk_kr = desc_text.split(' / ')[1]
    #                 desk_ru = desc_text.split(' / ')[0]
    #         else:
    #             desk_lt = desc_text.split(' / ')[2]
    #             desk_en = desc_text.split(' / ')[0]
    #             desk_ru = desc_text.split(' / ')[1]
    #         term = ' '.join(tx for tx in row_text.split())
    #         print(term)
    #         print(desk_lt)
    #         print(desk_kr)
    #         print(desk_en)
    #         print(desk_ru)
    q = 0
    p = 0
    r = 0
    for table in tables:
        for row in table.rows:
            if len(row.cells) == 2:
                terms = str(row.cells[0].text.strip())
                texts = str(row.cells[1].text.strip())
                if (" / " in terms or terms == terms.upper() and len(row.cells) == 2 and terms not in dont and terms
                        or "ТВт  ТВт" in terms):
                    q += 1
                    print(terms)
                    print(texts)
                elif "Ы " in terms and "E " in terms:

                    terms = terms.replace("\n", " ").replace("  ", " ")

                    p += 1
                    term_ru = terms[:terms.index('Ы ')]
                    term_en = terms[terms.index('E '):][2:]
                    term_uz = terms[terms.index('Ы '):terms.index('E ')][2:]
                    print(term_uz + "\n"+term_en + "\n"+term_ru + "\n\n")

    print(q)
    print(p)
    print(r)


def lugat4():
    file = "L 4 Электромагнитная-совместимость-радиотехнических-средств.-Термины-и-определения.docx"
    doc = Document(docx=file)

    tables = doc.tables
    p = 0
    q = 0
    r = 0
    for table in tables:
        for row in table.rows:
            terms = str(row.cells[1].text.strip())
            if "ru -" in terms or "en -" in terms:  # ('ru – ' in terms or 'ru -' in terms) and ("en - " in terms or 'en – ' in terms or 'en  - ' in terms)
                texts = str(row.cells[-1].text.strip())
                if len([item for item in texts.split(".\n\n") if item]) == 2:
                    text_uz = texts.split(".\n\n")[0].replace('\n', '').strip()
                    text_ru = texts.split(".\n\n")[1].replace('\n', '').strip()
                else:  # if len([item for item in texts.split(". \n\n") if item]) == 2:
                    text_uz = texts.split(". \n\n")[0].replace('\n', '').strip()
                    text_ru = texts.split(". \n\n")[1].replace('\n', '').strip()

                term_uz = terms[:terms.index('ru -')].strip()
                term_en = terms[terms.index('en -'):][3:].strip()
                term_ru = terms[terms.index('ru -'):terms.index('en -')][3:].strip()
                print(term_uz + "\n" + term_en + "\n" + term_ru + "\n\n")

                print(text_uz)
                print(text_ru + '\n\n')
    print(p)
    print(q)
    print(r)


def lugat5():
    file = "L 5 Словарь по электронной технике и радиоэлектронике.docx"
    doc = Document(docx=file)

    tables = doc.tables
    p = 0
    q = 0
    r = 0
    for table in tables:
        for row in table.rows:
            terms = str(row.cells[0].text.strip())
            if terms == str(row.cells[1].text.strip()):
                pass
            elif 'en - ' in terms and 'ru - ' in terms:
                q += 1
                texts = str(row.cells[-1].text.strip())
                text_uz = ''
                text_ru = ''
                if len([item for item in texts.split(".\n\n") if item]) == 2:
                    r += 1
                    text_uz = [item for item in texts.split(".\n\n") if item][0]
                    text_ru = [item for item in texts.split(".\n\n") if item][1]
                elif len([item for item in texts.split(". \n\n") if item]) == 2:
                    r += 1
                    text_uz = [item for item in texts.split(". \n\n") if item][0]
                    text_ru = [item for item in texts.split(". \n\n") if item][1]
                elif len([item for item in texts.split(".\n") if item]) == 2:
                    r += 1
                    text_uz = [item for item in texts.split(".\n") if item][0]
                    text_ru = [item for item in texts.split(".\n") if item][1]
                elif len([item for item in texts.split("\n\n") if item]) == 2:
                    r += 1
                    text_uz = [item for item in texts.split("\n\n") if item][0]
                    text_ru = [item for item in texts.split("\n\n") if item][1]
                else:
                    pass

                term_uz = terms[:terms.index('ru - ')].replace('\n', '').strip()
                term_en = terms[terms.index('en - '):].replace('\n', '').strip()
                term_ru = terms[terms.index('ru - '):terms.index('en - ')].replace('\n', '').strip()
                print(clean_text(term_uz) + "\n" + clean_text(term_en[5:].strip()) + "\n" + clean_text(
                    term_ru[5:].strip()))

                print(clean_text(text_uz))
                print(clean_text(text_ru) + '\n\n')
    print(q)
    print(r)
    print(p)


def lugat6():
    file = "L 6 Русско-узбекский толковый словарь терминов по телевидению.docx"
    doc = Document(docx=file)

    tables = doc.tables
    r = 0
    q = 0
    needs_row = []
    needs_desc = []
    for table in tables:
        rows = table.rows
        for row in rows:
            ter = row.cells[0].text.strip()
            if "ru - " in ter and "en - " in ter:
                needs_row.append(ter)
                needs_desc.append(row.cells[-1].text)

    check = False

    for term, desc in zip(needs_row, needs_desc):
        term_uz = clean_text(term[:term.index('ru - ')].replace('\n', '').strip())
        term_en = clean_text(term[term.index('en - '):].replace('\n', '').strip()[5:].strip())
        term_ru = clean_text(term[term.index('ru - '):term.index('en - ')].replace('\n', '').strip()[5:].strip())

        if len([item for item in desc.split(".\n\n") if item]) == 2:
            text_uz = [item for item in desc.split(".\n\n") if item][0].strip().replace('\n', '')
            text_ru = [item for item in desc.split(".\n\n") if item][1].strip().replace('\n', '')

        elif len([item for item in desc.split(". \n\n") if item]) == 2:
            text_uz = [item for item in desc.split(". \n\n") if item][0].strip().replace('\n', '')
            text_ru = [item for item in desc.split(". \n\n") if item][1].strip().replace('\n', '')

        else:  # len([item for item in desc.split(".\n") if item]) == 2
            text_uz = [item for item in desc.split(".\n") if item][0].strip().replace('\n', '')
            text_ru = [item for item in desc.split(".\n") if item][1].strip().replace('\n', '')

        print(term_uz)
        print(term_en)
        print(term_ru)
        print(text_uz)
        print(text_ru + '\n\n')


def lugat7():
    file = "L 7 Словарь сокращений по телекоммуникациям.docx"
    doc = Document(docx=file)

    row = 0
    row2 = 0
    tables = doc.tables
    for i in tables:
        row2 += len(i.rows)
        for r in i.rows:
            if r.cells[1].text == r.cells[2].text:
                row += 1
            elif r.cells[0] == "Қисқартмалар":
                pass
            else:
                shorts = r.cells[0].text.strip()
                uz = r.cells[3].text.strip()
                en = r.cells[1].text.strip()
                ru = r.cells[2].text.strip()


def lugat8():
    file = "L 8 Русско узбекский толковый словарь терминов по системам мобильной связи.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[1] == ''):
                pass
            elif "uz -" in row.cells[0].text or "ru -" in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[1].text)
                if 'uz -' in tex:
                    term_uz = tex[tex.index('uz -'):tex.index('en -')][4:].strip()
                    term_en = tex[tex.index('en -'):][5:].strip()
                    term_ru = tex[:tex.index('uz -')].strip()
                else:
                    term_uz = tex[tex.index('ru -'):tex.index('en -')][4:].strip()
                    term_en = tex[tex.index('en -'):][5:].strip()
                    term_ru = tex[:tex.index('ru -')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
            else:
                pass


def lugat9():
    file = "L 9 Русско узбекский толковый словарь терминов по линиям связи и системам передачи.docx"
    doc = Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[1] == ''):
                pass
            elif "uz - " in row.cells[0].text or "en - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
            else:
                pass


def lugat10():
    file = "L 10 Русско-узбекский толковый словарь терминов по электропитанию телекоммуникационных устройств.docx"
    doc = Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[1] == ''):
                pass
            elif "uz - " in row.cells[0].text and "en - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]

            else:
                pass


def lugat11():
    file = "L 11 Англо русско узбекский толковый словарь. Информационная технология. Операционные системы.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "ru - " in row.cells[0].text and "uz - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[-1].text)

                term_uz = tex[tex.index('uz - '):tex.index('ru - ')][5:].strip()
                term_en = tex[tex.index('ru - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]

            else:
                pass


def lugat12():
    file = "L 12 Словарь по информационной безопасности 1 издание.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[-1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                print(term_uz)
                print(term_en)
                print(term_ru)

            else:
                pass


def lugat13():
    file = "L 13 Русско узбекский толковый словарь терминов по радиотехнике.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[-1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]

            else:
                pass


def lugat14():
    file = "L 14 Русско узбекский толковый словарь терминов по коммутационным системам.docx"
    doc = Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]

            else:
                pass


def lugat15():
    file = "L 15 Русско узбекский толковый словарь терминов по системам беспроводного доступа.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
            else:
                pass


def lugat16():
    file = "L 16 Русско узбекский толковый словарь терминов по измерениям в телекоммуникации.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
            else:
                pass


def lugat17():
    file = "L 17 Русско узбекский толковый словарь терминов по радиочастотному спектру_ радиоэлектронным средств.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat18():
    file = "L 18 Русско-узбекский толковый словарь терминов по сетям передачи данных.docx"
    doc = Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat19():
    file = "L 19 Русско-узбекский толковый словарь терминов по программированию.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat20():
    file = "L 20 Русско узбекский толковый словарь терминов по спутниковой связи.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat21():
    file = "L 21 Русско-узбекский толковый словарь терминов по телекоммуникационным услугам.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat22():
    file = "L 22 Русско-узбекский толковый словарь терминов по вычислительной технике.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat23():
    file = "L 23 Русско-узбекский толковый словарь терминов по электронному документообороту.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat24():
    file = "L 24 Русско-узбекский толковый словарь терминов по оптоэлектронике.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat25():
    file = "L 25 Русско-узбекский толковый словарь терминов по радиорелейным системам.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat26():
    file = "L 26 Русско-узбекский толковый словарь по менеджменту качества.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:
                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat27():
    file = "L 27 Англо-русско-узбекский словарь сокращений терминов по телекоммуникациям.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if (row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == '')
                    or row.cells[-1].text == "Ўзбек тилидаги тўлиқ номланиши"):
                pass
            elif len(row.cells) >= 5:
                if len(row.cells) == 5:
                    tex = row.cells[0].text.strip()
                    desc = row.cells

                    term = tex
                    desc_en = desc[1].text.strip()
                    desc_uz = desc[3].text.strip()
                    desc_uz_kr = desc[-1].text.strip()
                    desc_ru = desc[2].text.strip()
                else:
                    texts = []
                    for el in row.cells:
                        if el.text.strip() not in texts:
                            texts.append(el.text.strip())
                    if len(texts) < 5:
                        texts.append(texts[2])
                    term = texts[0]
                    desc_en = texts[1]
                    desc_uz = texts[2]
                    desc_uz_kr = texts[3]
                    desc_ru = texts[4]


def lugat28():
    file = "L 28 Русско-узбекский толковый словарь терминов по телевидению.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat29():
    file = "L 29 Русско-узбекский толковый словарь терминов по электронной технике и радиоэлектронике.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:
                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat30():
    file = "L 30 Русско-узбекский толковый словарь терминов по современным компьютерным технологиям.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat31():
    file = "L 31 Русско-узбекский толковый словарь терминов по информационной безопасности.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat32():
    file = "L 32 Русско-узбекский толковый словарь терминов по нанотехнологиям.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]




def lugat34():
    file = "L 34 Русско-узбекский толковый словарь терминов по распространению радиоволн и антенно-фидерным устройствам. 1 часть.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat35():
    file = "L 35 Русско-узбекский толковый словарь терминов по мобильной и фиксированной беспроводной связи.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat36():
    file = "L 36 Русско-узб. тол. словарь терминов операционных систем и компьютерных сетей.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat37():
    file = "L 37 Русско узб. тол. словарь по маркетингу и экономикии.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat38():
    file = "L 38 Словарь по вычислительной технике.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat39():
    file = "L 39 Русско-узбекский толковый словарь терминов в системе ''Электронное правительство''.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]


def lugat40():
    file = "L 40 Словарь специальных терминов по беспилотной авиационной технике.docx"
    doc = Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]

                    if len(term_uz.split('\n ')) == 2:
                        term_uz_ln = term_uz.split('\n ')[0]
                        term_uz_kr = term_uz.split('\n ')[1]